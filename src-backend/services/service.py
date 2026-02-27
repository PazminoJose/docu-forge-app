from sanic import Request
from sanic.response import json, file, raw
import pandas as pd
from docx import Document
import os
import re
from datetime import datetime
from openpyxl.utils import column_index_from_string, get_column_letter
from openpyxl import load_workbook
import numpy as np
import asyncio


async def get_docx_file(request: Request):
    try:
        # Obtenemos el path de los parámetros de la URL
        path = request.args.get("path")

        if not os.path.exists(path):
            return json({"error": "File not found"}, status=404)

        with open(path, "rb") as f:
            file_bytes = f.read()

        # Usamos raw para enviar bytes puros
        # El content_type es vital para que el navegador no intente leerlo como texto
        return raw(
            file_bytes, 
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print(f"ERROR AL SERVIR ARCHIVO: {e}")
        return json({"error": "Error interno al leer el archivo"}, status=500)


async def get_docx_fields(request: Request):
    try:
        path = request.args.get("path", "")
        
        if not path or not os.path.exists(path):
            return json({"error": "Invalid path or file not found"}, status=400)

        # Exec the extraction in a separate thread to avoid blocking Sanic's event loop
        fields = await asyncio.to_thread(_extract_fields, path)

        return json({"fields": fields})
        
    except Exception as e:
        print(f"DEBUG ERROR: {e}")
        return json({"error": str(e)}, status=500)
    
def _extract_fields(path: str) -> list[str]:
    doc = Document(path)
    placeholders = set()
    regex = r"\$\{([^}]+)\}"

    # Search for matches in a list of paragraphs 
    def find_in_paragraphs(paragraphs):
        for p in paragraphs:
            matches = re.findall(regex, p.text)
            for m in matches:
                placeholders.add(m)

    # 1. Main body paragraphs
    find_in_paragraphs(doc.paragraphs)

    # 2. Tables paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                find_in_paragraphs(cell.paragraphs)
                
    # 3. Headers and footers paragraphs
    for section in doc.sections:
        find_in_paragraphs(section.header.paragraphs)
        find_in_paragraphs(section.footer.paragraphs)

    return list(placeholders)

# async def get_mapper_data(request: Request):
#     try:
#         path = request.args.get("path", "")
#         start = int(request.args.get("start", 0))
#         limit = int(request.args.get("limit", 10))
        
#         if not path or not os.path.exists(path):
#             return json({"error": "Invalid path"}, status=400)

#         # Leer sin encabezados
#         df = pd.read_excel(path, engine="openpyxl", header=None)

#         # Reemplazar NaN por None (para que sea null en JSON)
#         df = df.replace({np.nan: None})

#         # Tomar el slice (paginación)
#         df_slice = df.iloc[start : start + limit]

#         # Convertir a lista de listas
#         raw_rows = df_slice.values.tolist()

#         # LIMPIEZA MANUAL: Convertir cualquier objeto datetime a string
#         clean_data = []
#         for row in raw_rows:
#             clean_row = []
#             for cell in row:
#                 # 1. Manejar fechas
#                 if isinstance(cell, (datetime, pd.Timestamp)):
#                     clean_row.append(cell.strftime("%Y-%m-%d %H:%M:%S"))
#                 # 2. Manejar nulos (NaN / None)
#                 elif pd.isna(cell): 
#                     clean_row.append("NA") # O "" si prefieres string vacío en lugar de null
#                 else:
#                     clean_row.append(cell)
#             clean_data.append(clean_row)


#         return json({"data": clean_data})
        
#     except Exception as e:
#         # Imprimir el error en consola para debugging
#         print(f"DEBUG ERROR: {e}")
#         return json({"error": str(e)}, status=500)
async def get_mapper_data(request: Request):
    try:
        path = request.args.get("path", "")
        start = int(request.args.get("start", 0))
        limit = int(request.args.get("limit", 10))
        
        if not path or not os.path.exists(path):
            return json({"error": "Invalid path"}, status=400)

        # 1. Leer TODAS las hojas (sheet_name=None devuelve un dict)
        all_sheets_df = pd.read_excel(path, engine="openpyxl", header=None, sheet_name=None)

        final_response = []

        for sheet_name, df in all_sheets_df.items():
            # Reemplazar NaN por None
            df = df.replace({np.nan: None})

            # Tomar el slice (paginación)
            df_slice = df.iloc[start : start + limit]

            sheet_data = []
            for _, row in df_slice.iterrows():
                clean_row = []
                for idx, cell in enumerate(row):
                    # Identificar la letra de la columna (A, B, C...)
                    col_letter = get_column_letter(idx + 1)
                    
                    # Formatear el valor
                    value = ""
                    if isinstance(cell, (datetime, pd.Timestamp)):
                        value = cell.strftime("%Y-%m-%d %H:%M:%S")
                    elif cell is None or pd.isna(cell):
                        value = "" # O "NA" según prefieras
                    else:
                        value = cell

                    # Guardamos como objeto con metadatos de columna
                    clean_row.append({
                        "column": col_letter,
                        "value": value,
                    })
                sheet_data.append(clean_row)

            # Estructura solicitada
            final_response.append({
                "sheet": sheet_name,
                "data": sheet_data
            })

        return json(final_response)
        
    except Exception as e:
        print(f"DEBUG ERROR: {e}")
        return json({"error": str(e)}, status=500)

async def get_mapper_range(request: Request):
    try:
        path = request.args.get("path", "")
        
        if not path or not os.path.exists(path):
            return json({"error": "Invalid path"}, status=400)

        all_sheets = pd.read_excel(path, engine="openpyxl", header=None, sheet_name=None)

        ranges_response = []

        for sheet_name, df in all_sheets.items():
            total_rows = len(df)
            
            ranges_response.append({
                "sheet": sheet_name,
                "range": {
                    "from": 1 if total_rows > 0 else 0,
                    "to": total_rows
                }
            })

        return json(ranges_response)

    except Exception as e:
        print(f"DEBUG ERROR: {e}")
        return json({"error": str(e)}, status=500)   
    
# async def get_mapper_range(request: Request):
#     try:
#         path = request.args.get("path", "")
        
#         if not path or not os.path.exists(path):
#             return json({"error": "Invalid path"}, status=400)

#         df = pd.read_excel(path, engine="openpyxl", header=None)
#         total_rows = len(df)

#         data_range = {
#             "from": 1,
#             "to": total_rows
#         }

#         return json({"range": data_range})
#     except Exception as e:
#         print(f"DEBUG ERROR: {e}")
#         return json({"error": str(e)}, status=500)


async def get_unique_combinations(request):
    try:
        data = request.json
        xlsx_path = data.get("path")
        filters = data.get("filters", []) 

        if not xlsx_path or not os.path.exists(xlsx_path):
            return json({"error": "Archivo no encontrado"}, status=404)

        if not filters:
            return json({"error": "No se enviaron filtros"}, status=400)

        # 1. Cargar el Excel (solo valores para velocidad)
        wb = load_workbook(xlsx_path, data_only=True, read_only=True)
        sheet = wb.active

        # 2. Convertir letras de columna (A, B...) a índices numéricos (1, 2...)
        col_indices = []
        for f in filters:
            try:
                col_indices.append(column_index_from_string(f['column']))
            except ValueError:
                return json({"error": f"Columna inválida: {f['column']}"}, status=400)

        # 3. Extraer combinaciones únicas
        unique_combinations = set()
        
        # Iteramos las filas saltando el header (asumiendo que hay uno)
        # Si no hay header, quita el min_row=2
        for row in sheet.iter_rows(min_row=2, values_only=True):
            # Extraemos solo las celdas de las columnas solicitadas
            # Nota: row está indexado en 0, col_indices en 1
            combination = tuple(row[idx - 1] for idx in col_indices)
            
            # Solo agregar si al menos uno de los valores de la combinación no es None
            if any(val is not None for val in combination):
                unique_combinations.add(combination)

        # 4. Formatear la respuesta
        # Convertimos el set de tuplas a una lista de diccionarios para el Frontend
        result = []
        for combo in unique_combinations:
            combination_group = []
            for i, f in enumerate(filters):
                combination_group.append({
                    "column": f['column'],
                    "value": combo[i] if combo[i] is not None else ""
                })
            
            result.append(combination_group)
            
        return json({
            "combinations": result
        })

    except Exception as e:
        return json({"status": "error", "message": str(e)}, status=500)