from sanic import Websocket
import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string
from docx import Document
import os
from datetime import datetime
import json as py_json
import asyncio
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

async def process_docx(ws):
    try:
        # 1. Recepción de configuración inicial
        raw_msg = await ws.recv()
        data = py_json.loads(raw_msg)
        
        if data.get("action") != "start":
            return

        # Parámetros base
        docx_path = data.get("docx_path")
        xlsx_path = data.get("xlsx_path")
        output_root = data.get("output_folder")
        fields = data.get("fields", [])
        skip_header = data.get("skip_header", False)
        apply_all = data.get("apply_to_all_sheets", False)
        merge_enabled = data.get("merge_generated_files", False)

        if not os.path.exists(xlsx_path):
            await ws.send(py_json.dumps({"status": "error", "message": "Archivo Excel no encontrado"}))
            return

        # Cargar libro de Excel
        wb = load_workbook(xlsx_path, data_only=True)
        
        # --- 2. Definir Hojas a Procesar ---
        sheets_to_process = []
        if apply_all:
            for name in wb.sheetnames:
                sheet = wb[name]
                start = 2 if skip_header else 1
                sheets_to_process.append({
                    "name": name, 
                    "start": start, 
                    "end": sheet.max_row
                })
        else:
            range_config = data.get("range", {})
            start = int(range_config.get("from", 1))
            end = int(range_config.get("to", 1))
            if skip_header: start += 1
            
            target_sheet = next((f.get("mappedToSheet") for f in fields if f.get("mappedToSheet")), wb.active.title)
            sheets_to_process.append({"name": target_sheet, "start": start, "end": end})

        # Cálculo de progreso global
        total_work = sum(max(0, s["end"] - s["start"] + 1) for s in sheets_to_process)
        current_global_idx = 0

        # --- 3. Bucle Principal de Hojas ---
        for s_info in sheets_to_process:
            sheet_name = s_info["name"]
            if sheet_name not in wb.sheetnames: continue
            
            sheet = wb[sheet_name]
            generated_paths_in_folder = []
            
            current_output_folder = output_root
            if apply_all:
                clean_sheet_name = "".join(c for c in sheet_name if c.isalnum() or c in (' ', '_', '-')).strip()
                current_output_folder = os.path.join(output_root, clean_sheet_name)
            
            if not os.path.exists(current_output_folder):
                os.makedirs(current_output_folder, exist_ok=True)

            # --- 4. Bucle de Filas ---
            for row_idx in range(s_info["start"], s_info["end"] + 1):
                try:
                    client_msg = await asyncio.wait_for(ws.recv(), timeout=0.0001)
                    if py_json.loads(client_msg).get("action") == "cancel":
                        await ws.send(py_json.dumps({"status": "cancelled"}))
                        return
                except: pass

                doc = Document(docx_path)
                was_modified = False
                custom_filename = None

                for field in fields:
                    identifier = f"${{{field['identifier']}}}"
                    col_letter = field.get('mappedToColumn')
                    
                    if col_letter:
                        col_idx = column_index_from_string(col_letter)
                        if not apply_all and field.get("mappedToSheet") and field.get("mappedToSheet") != sheet_name:
                            source_sheet = wb[field.get("mappedToSheet")]
                            cell_value = source_sheet.cell(row=row_idx, column=col_idx).value if row_idx <= source_sheet.max_row else None
                        else:
                            cell_value = sheet.cell(row=row_idx, column=col_idx).value if row_idx <= sheet.max_row else None
                    else:
                        cell_value = field.get('value', "")

                    if isinstance(cell_value, (datetime, pd.Timestamp)):
                        replacement = cell_value.strftime("%Y-%m-%d")
                    else:
                        replacement = str(cell_value) if cell_value is not None else ""

                    if field.get('useAsName', False) and replacement.strip():
                        custom_filename = "".join(c for c in replacement if c.isalnum() or c in (' ', '_', '-')).strip()

                    if _replace_cleanly(doc, identifier, replacement):
                        was_modified = True

                if was_modified:
                    fname = f"{custom_filename}.docx" if custom_filename else f"{sheet_name}_Fila_{row_idx}.docx"
                    save_path = os.path.join(current_output_folder, fname)
                    doc.save(save_path)
                    generated_paths_in_folder.append(save_path)

                current_global_idx += 1
                percent = int((current_global_idx / total_work) * 100) if total_work > 0 else 100
                await ws.send(py_json.dumps({
                    "status": "progress",
                    "percent": percent,
                    "current": current_global_idx,
                    "total": total_work,
                    "sheet": sheet_name
                }))
                await asyncio.sleep(0.01)

            # --- 5. Lógica de MERGE MEJORADA ---
            if merge_enabled and generated_paths_in_folder:
                await ws.send(py_json.dumps({"status": "progress_info", "message": f"Consolidando archivos de {sheet_name}..."}))
                
                consolidated_path = os.path.join(current_output_folder, f"Consolidado_{sheet_name}.docx")
                _merge_documents(generated_paths_in_folder, consolidated_path)

                # Eliminar archivos individuales, dejando solo el consolidado
                for p_path in generated_paths_in_folder:
                    try:
                        if os.path.exists(p_path) and os.path.abspath(p_path) != os.path.abspath(consolidated_path):
                            os.remove(p_path)
                    except OSError:
                        pass

        await ws.send(py_json.dumps({"status": "completed"}))

    except Exception as e:
        import traceback
        print(f"CRITICAL ERROR:\n{traceback.format_exc()}")
        await ws.send(py_json.dumps({"status": "error", "message": str(e)}))

def _merge_documents(doc_paths, output_path):
    """
    Une múltiples .docx en uno solo. Cada documento fuente comienza
    en su propia página (salto de sección). Se preservan imágenes,
    tablas y formato. Se eliminan páginas vacías resultantes.
    """
    if not doc_paths:
        return

    # Usar el primer documento como base (preserva estilos, headers, etc.)
    merged = Document(doc_paths[0])

    for doc_path in doc_paths[1:]:
        sub_doc = Document(doc_path)

        # Insertar un salto de sección de página nueva en el documento merged
        # para que el contenido del siguiente doc empiece en página propia
        _add_page_break_section(merged)

        # Copiar todas las relaciones (imágenes, etc.) del sub_doc al merged
        rel_mapping = _copy_relationships(sub_doc, merged)

        # Copiar cada elemento del body del sub_doc
        for element in sub_doc.element.body:
            # No copiar el sectPr final del sub_doc (es su config de sección)
            if element.tag == qn('w:sectPr'):
                continue
            
            new_element = copy.deepcopy(element)
            
            # Reasignar referencias a relaciones (imágenes embebidas)
            _remap_relationships(new_element, rel_mapping)
            
            merged.element.body.append(new_element)

    # Eliminar páginas/párrafos vacíos resultantes
    _remove_empty_pages(merged)

    merged.save(output_path)

def _add_page_break_section(doc):
    """
    Añade un salto de sección 'nextPage' al final del documento,
    forzando que lo siguiente empiece en página nueva.
    Copia también las referencias a headers/footers para preservar
    encabezados e imágenes de encabezado.
    """
    # Obtener el último párrafo o crear uno
    last_p = doc.element.body.findall(qn('w:p'))
    if last_p is not None and len(last_p) > 0:
        target_p = last_p[-1]
    else:
        target_p = OxmlElement('w:p')
        doc.element.body.append(target_p)

    # Crear propiedades del párrafo con sección nextPage
    pPr = target_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        target_p.insert(0, pPr)

    sectPr = OxmlElement('w:sectPr')
    sect_type = OxmlElement('w:type')
    sect_type.set(qn('w:val'), 'nextPage')
    sectPr.append(sect_type)

    # Copiar propiedades de la sección principal del documento
    body_sectPr = doc.element.body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        # Tamaño de página
        pgSz = body_sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            sectPr.append(copy.deepcopy(pgSz))
        # Márgenes
        pgMar = body_sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            sectPr.append(copy.deepcopy(pgMar))
        # Columnas
        cols = body_sectPr.find(qn('w:cols'))
        if cols is not None:
            sectPr.append(copy.deepcopy(cols))
        # Bordes de página
        pgBorders = body_sectPr.find(qn('w:pgBorders'))
        if pgBorders is not None:
            sectPr.append(copy.deepcopy(pgBorders))

        # --- CLAVE: Copiar referencias a headers y footers ---
        for headerRef in body_sectPr.findall(qn('w:headerReference')):
            sectPr.append(copy.deepcopy(headerRef))
        for footerRef in body_sectPr.findall(qn('w:footerReference')):
            sectPr.append(copy.deepcopy(footerRef))

        # Copiar titlePg si existe (indica que la primera página tiene header distinto)
        titlePg = body_sectPr.find(qn('w:titlePg'))
        if titlePg is not None:
            sectPr.append(copy.deepcopy(titlePg))

    pPr.append(sectPr)

def _copy_relationships(source_doc, target_doc):
    """
    Copia las relaciones (imágenes, etc.) de source a target.
    Retorna un dict {old_rId: new_rId} para reasignar referencias.
    """
    rel_mapping = {}
    
    source_part = source_doc.part
    target_part = target_doc.part

    for rel in source_part.rels.values():
        # Solo copiar relaciones de imágenes y otros embebidos
        if "image" in rel.reltype or "oleObject" in rel.reltype or "chart" in rel.reltype:
            try:
                # Obtener el blob de la imagen
                image_part = rel.target_part
                
                # Crear nueva relación en el documento destino
                new_rel = target_part.relate_to(image_part, rel.reltype)
                rel_mapping[rel.rId] = new_rel
            except Exception:
                # Si falla, mantener el rId original
                rel_mapping[rel.rId] = rel.rId

    return rel_mapping

def _remap_relationships(element, rel_mapping):
    """
    Recorre un elemento XML y reemplaza los rId viejos por los nuevos
    según el mapping. Esto asegura que las imágenes apunten correctamente.
    """
    if not rel_mapping:
        return
    
    # Buscar todos los atributos que contengan rId (blip, embed, etc.)
    # Namespace de relaciones
    r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    
    for child in element.iter():
        # Atributo r:embed (imágenes inline)
        embed = child.get(qn('r:embed'))
        if embed and embed in rel_mapping:
            child.set(qn('r:embed'), rel_mapping[embed])
        
        # Atributo r:link (imágenes enlazadas)
        link = child.get(qn('r:link'))
        if link and link in rel_mapping:
            child.set(qn('r:link'), rel_mapping[link])
        
        # Atributo r:id (objetos OLE, charts)
        rid = child.get(qn('r:id'))
        if rid and rid in rel_mapping:
            child.set(qn('r:id'), rel_mapping[rid])

def _remove_empty_pages(doc):
    """
    Elimina párrafos vacíos que generarían páginas en blanco.
    """
    body = doc.element.body
    paragraphs_to_remove = []

    # Namespace de mc (Markup Compatibility)
    MC_NS = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    MC_ALTERNATE = f'{{{MC_NS}}}AlternateContent'

    all_elements = list(body)
    
    for i, element in enumerate(all_elements):
        if element.tag != qn('w:p'):
            continue

        # Obtener texto del párrafo
        text = ''.join(node.text or '' for node in element.iter(qn('w:t'))).strip()
        
        # Verificar si contiene contenido visual (imágenes, drawings, etc.)
        has_drawing = bool(element.findall('.//' + qn('w:drawing')))
        has_pict = bool(element.findall('.//' + qn('w:pict')))
        has_mc = bool(element.findall('.//' + MC_ALTERNATE))
        has_visual_content = has_drawing or has_pict or has_mc

        if has_visual_content:
            continue

async def process_multiple_docx(ws: Websocket):
    try:
        raw_msg = await ws.recv()
        data = py_json.loads(raw_msg)
        
        if data.get("action") == "start":
            # 1. Extracción de Parámetros
            xlsx_path = data.get("xlsx_path")
            output_root = data.get("output_folder")
            fields = data.get("fields", [])
            filters_config = data.get("filters", [])
            template_mappings = data.get("templateMapping", [])
            range_config = data.get("range", {})
            
            start_row = int(range_config.get("from", 1))
            end_row = int(range_config.get("to", 1))
            if data.get("skipHeader", True): start_row += 1

            # 2. Preparación de Mapeos Rápidos
            # Creamos un diccionario: {(val_col_B, val_col_C): "path/to/template.docx"}
            mapping_dict = {}
            for item in template_mappings:
                # Ordenamos la combinación por letra de columna para asegurar consistencia
                combo_key = tuple(c['value'] for c in item['combination'])
                mapping_dict[combo_key] = item['templatePath']

            # Identificar qué columna dicta la subcarpeta
            folder_filter = next((f for f in filters_config if f.get("useAsFolderName")), None)
            folder_col_letter = folder_filter["column"] if folder_filter else None

            # 3. Carga de Excel
            wb = load_workbook(xlsx_path, data_only=True)
            sheet = wb.active
            total_rows = end_row - start_row + 1

            # 4. Loop de Procesamiento
            for i, row_idx in enumerate(range(start_row, end_row + 1)):
                # --- A. Verificación de Cancelación ---
                try:
                    client_msg = await asyncio.wait_for(ws.recv(), timeout=0.0001)
                    if py_json.loads(client_msg).get("action") == "cancel":
                        await ws.send(py_json.dumps({"status": "cancelled"}))
                        return
                except: pass

                # --- B. Determinar Combinación de la Fila Actual ---
                current_row_combo = []
                for f in filters_config:
                    col_idx = column_index_from_string(f["column"])
                    val = sheet.cell(row=row_idx, column=col_idx).value
                    current_row_combo.append(str(val) if val is not None else "")
                
                row_combo_tuple = tuple(current_row_combo)

                # --- C. ¿Esta fila tiene una plantilla asignada? ---
                if row_combo_tuple not in mapping_dict:
                    continue # Saltar fila si no coincide con ninguna combinación

                target_template = mapping_dict[row_combo_tuple]
                
                # --- D. Determinar Carpeta de Destino ---
                final_output_folder = output_root
                if folder_col_letter:
                    subfolder_name = str(sheet.cell(row=row_idx, column=column_index_from_string(folder_col_letter)).value or "Sin_Nombre")
                    # Limpiar nombre de carpeta de caracteres inválidos
                    subfolder_name = "".join(c for c in subfolder_name if c.isalnum() or c in (' ', '_', '-')).strip()
                    final_output_folder = os.path.join(output_root, subfolder_name)
                
                if not os.path.exists(final_output_folder):
                    os.makedirs(final_output_folder, exist_ok=True)

                # --- E. Procesar el Documento (Reutilizando tu lógica) ---
                doc = Document(target_template)
                was_modified = False
                custom_filename = None

                for field in fields:
                    identifier = f"${{{field['identifier']}}}"
                    col_letter = field.get('mappedToColumn')
                    
                    if col_letter:
                        val = sheet.cell(row=row_idx, column=column_index_from_string(col_letter)).value
                    else:
                        val = field.get('value', "")

                    # Formateo
                    if isinstance(val, (datetime)):
                        replacement = val.strftime("%Y-%m-%d")
                    else:
                        replacement = str(val) if val is not None else ""

                    if field.get('useAsName') and replacement.strip():
                        custom_filename = "".join(c for c in replacement if c.isalnum() or c in (' ', '_', '-')).strip()

                    if _replace_cleanly(doc, identifier, replacement):
                        was_modified = True

                # --- F. Guardar ---
                if was_modified:
                    fname = f"{custom_filename}.docx" if custom_filename else f"Fila_{row_idx}.docx"
                    doc.save(os.path.join(final_output_folder, fname))

                # --- G. Reportar Progreso ---
                percent = int(((i + 1) / total_rows) * 100)
                await ws.send(py_json.dumps({
                    "status": "progress",
                    "percent": percent,
                    "current": i + 1,
                    "total": total_rows
                }))
                await asyncio.sleep(0.01)

            await ws.send(py_json.dumps({"status": "completed"}))

    except Exception as e:
        print(f"ERROR: {e}")
        await ws.send(py_json.dumps({"status": "error", "message": str(e)}))

def _replace_cleanly(doc, old_text, new_text):
    found = False
    # Función auxiliar para procesar párrafos de forma uniforme
    def process_p(p):
        nonlocal found
        if old_text in p.text:
            _combine_runs(p)  # Paso 1: Sanear el párrafo
            for run in p.runs: # Paso 2: Reemplazo limpio
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    found = True

    for p in doc.paragraphs:
        process_p(p)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_p(p)
    return found

def _combine_runs(paragraph):
    if len(paragraph.runs) < 2:
        return
    # iterar de atrás hacia adelante para poder eliminar runs sin romper el índice
    for i in range(len(paragraph.runs) - 1, 0, -1):
        r1 = paragraph.runs[i-1]
        r2 = paragraph.runs[i]
        # Si tienen el mismo formato se combinan
        if r1.bold == r2.bold and r1.italic == r2.italic and r1.underline == r2.underline and r1.font.name == r2.font.name:
            r1.text += r2.text
            # Eliminar el run r2 (técnicamente se elimina el elemento XML)
            r2._element.getparent().remove(r2._element)

