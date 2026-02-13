from sanic import Sanic
from sanic.response import json
from pandas import read_excel
import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string
from docx import Document
import os
from datetime import datetime
import numpy as np

app = Sanic("PythonBackend")

@app.get("/show-data")
async def data(request):
    try:
        path = request.args.get("path", "")
        start = int(request.args.get("start", 0))
        limit = int(request.args.get("limit", 10))
        
        if not path or not os.path.exists(path):
            return json({"error": "Invalid path"}, status=400)

        # Leer sin encabezados
        df = pd.read_excel(path, engine="openpyxl", header=None)

        # Reemplazar NaN por None (para que sea null en JSON)
        df = df.replace({np.nan: None})

        # Tomar el slice (paginación)
        df_slice = df.iloc[start : start + limit]

        # Convertir a lista de listas
        raw_rows = df_slice.values.tolist()

        # LIMPIEZA MANUAL: Convertir cualquier objeto datetime a string
        clean_data = []
        for row in raw_rows:
            clean_row = []
            for cell in row:
                # 1. Manejar fechas
                if isinstance(cell, (datetime, pd.Timestamp)):
                    clean_row.append(cell.strftime("%Y-%m-%d %H:%M:%S"))
                # 2. Manejar nulos (NaN / None)
                elif pd.isna(cell): 
                    clean_row.append("NA") # O "" si prefieres string vacío en lugar de null
                else:
                    clean_row.append(cell)
            clean_data.append(clean_row)

        data_range = {
            "from": 1,
            "to": len(df)
        }

        return json({"data": clean_data, "range": data_range})
        
    except Exception as e:
        # Imprimir el error en consola para debugging
        print(f"DEBUG ERROR: {e}")
        return json({"error": str(e)}, status=500)
    
@app.post("/process-template-docx")
def process_template_docx(request):
    # --- 1. Obtención de Parámetros Raíz ---
    docx_path = request.json.get("docx_path")
    xlsx_path = request.json.get("xlsx_path")
    output_folder = request.json.get("output_folder")
    fields = request.json.get("fields")
    skip_header = request.json.get("skip_header")
    range = request.json.get("range")
    start_row = int(range.get("from"))
    end_row = int(range.get("to"))
    
    # Ajuste por encabezado
    if skip_header:
        start_row += 1
    try:
        # data_only=True es vital para traer valores y no fórmulas
        wb = load_workbook(xlsx_path, data_only=True)
        sheet = wb.active
        
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
    except Exception as e:
        return json({"error": f"Error al acceder a archivos: {str(e)}"}, status=400)

    # --- 3. Procesamiento por Filas ---
    for row_idx in range(start_row, end_row + 1):
        doc = Document(docx_path)
        was_modified = False
        custom_filename = None
        
        for field in fields:
            identifier = f"${{{field['identifier']}}}"
            col_letter = field.get('mappedToColumn', 'A')
            col_idx = column_index_from_string(col_letter)
            
            cell_value = sheet.cell(row=row_idx, column=col_idx).value
            
            # Formateo de reemplazo (Strings, fechas y nulos)
            if isinstance(cell_value, (datetime, pd.Timestamp)):
                replacement = cell_value.strftime("%Y-%m-%d")
            else:
                replacement = str(cell_value) if cell_value is not None else ""
            
            # Nombre de archivo dinámico
            if field.get('useAsName', False) and replacement.strip():
                # Sanitización de caracteres para Windows/Linux
                custom_filename = "".join(c for c in replacement if c.isalnum() or c in (' ', '_', '-')).strip()

            # Reemplazo con protección de Runs (Mantiene Negritas/Cursivas)
            if _replace_cleanly(doc, identifier, replacement):
                was_modified = True

        # --- 4. Guardado ---
        if was_modified:
            final_name = f"{custom_filename}.docx" if custom_filename else f"Resultado_Fila_{row_idx}.docx"
            save_path = os.path.join(output_folder, final_name)
            doc.save(save_path)
            
    return json({"status": "success", "message": "Procesamiento completado"})
    
def _replace_cleanly(doc, old_text, new_text):
    found = False
    
    # Función auxiliar para procesar párrafos de forma uniforme
    def process_p(p):
        nonlocal found
        if old_text in p.text:
            _combine_runs(p)  # Paso 1: Sanar el párrafo
            for run in p.runs: # Paso 2: Reemplazo limpio
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    found = True

    for p in doc.paragraphs:
        process_p(p)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_p(p)
    return found

def _combine_runs(paragraph):
    if len(paragraph.runs) < 2:
        return
    # Iteramos de atrás hacia adelante para poder eliminar runs sin romper el índice
    for i in range(len(paragraph.runs) - 1, 0, -1):
        r1 = paragraph.runs[i-1]
        r2 = paragraph.runs[i]
        # Si tienen el mismo formato, los combinamos
        if r1.bold == r2.bold and r1.italic == r2.italic and r1.underline == r2.underline and r1.font.name == r2.font.name:
            r1.text += r2.text
            # Eliminar el run r2 (técnicamente eliminamos el elemento XML)
            r2._element.getparent().remove(r2._element)


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=8000, debug=False, fast=False, workers=1)
